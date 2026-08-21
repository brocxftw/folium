"""Local text extraction without LLMs."""

from __future__ import annotations

import logging
import re
import tempfile
from collections.abc import Callable
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import pymupdf
from docx import Document as DocxDocument
from PIL import Image

from folium.core.config import Settings, get_settings
from folium.core.exceptions import ValidationError
from folium.ocr.paddle_engine import ocr_image, paddle_ocr_available
from folium.ocr.subprocess_client import OcrSubprocessError, run_ocr_subprocess

logger = logging.getLogger(__name__)

_MIN_CHARS_PER_PAGE = 32
_DEFAULT_OCR_DPI = 150

OnOcrProgress = Callable[[int, int], None]
OnOcrPage = Callable[[int, str], None]


@dataclass(frozen=True)
class ExtractedPage:
    page_number: int
    text: str


@dataclass
class ExtractedDocument:
    pages: list[ExtractedPage] = field(default_factory=list)
    page_count: int = 0
    language: str | None = None
    method: str = "unknown"

    def to_page_dicts(self) -> list[dict[str, Any]]:
        return [{"page_number": p.page_number, "text": p.text} for p in self.pages]


def extract_document(
    path: Path,
    mime_type: str,
    *,
    settings: Settings | None = None,
    language: str | None = None,
    allow_ocr: bool = True,
    force_ocr: bool = False,
    on_ocr_progress: OnOcrProgress | None = None,
    on_ocr_page: OnOcrPage | None = None,
) -> ExtractedDocument:
    """Extract text from a supported document on disk.

    For PDFs, OCR is optional:
    - ``allow_ocr=False``: native text only (fast pre-flight extract)
    - ``force_ocr=True``: always run OCR (dedicated OCR job)
    - default: OCR when native text looks too thin

    When ``on_ocr_page`` is provided, each OCR page is reported as it completes
    (for incremental persistence). Returned ``pages`` still includes all pages
    unless the caller only needs the stream.
    """
    settings = settings or get_settings()
    lang = language or settings.ocr_language
    path = path.resolve()
    if not path.is_file():
        raise ValidationError(f"File not found: {path}")

    mime = mime_type.lower()
    if mime == "application/pdf":
        return _extract_pdf(
            path,
            settings=settings,
            language=lang,
            allow_ocr=allow_ocr,
            force_ocr=force_ocr,
            on_ocr_progress=on_ocr_progress,
            on_ocr_page=on_ocr_page,
        )
    if mime in {"image/png", "image/jpeg"}:
        return _extract_image(
            path,
            settings=settings,
            language=lang,
            on_ocr_progress=on_ocr_progress,
            on_ocr_page=on_ocr_page,
        )
    if mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_docx(path)
    if mime in {"text/plain", "text/markdown"}:
        return _extract_text_file(path)
    raise ValidationError(f"Unsupported MIME type for extraction: {mime_type}")


def _extract_text_file(path: Path) -> ExtractedDocument:
    text = path.read_text(encoding="utf-8")
    pages = _split_text_pages(text)
    return ExtractedDocument(
        pages=pages,
        page_count=len(pages),
        language=None,
        method="utf8",
    )


def _split_text_pages(text: str) -> list[ExtractedPage]:
    if "\f" in text:
        parts = text.split("\f")
    else:
        parts = [text]
    return [
        ExtractedPage(page_number=index + 1, text=part.strip())
        for index, part in enumerate(parts)
        if part.strip()
    ] or [ExtractedPage(page_number=1, text="")]


def _extract_docx(path: Path) -> ExtractedDocument:
    doc = DocxDocument(str(path))
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        line = para.text.strip()
        if line:
            paragraphs.append(line)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    text = "\n\n".join(paragraphs)
    return ExtractedDocument(
        pages=[ExtractedPage(page_number=1, text=text)],
        page_count=1,
        language=None,
        method="docx",
    )


def _extract_pdf(
    path: Path,
    *,
    settings: Settings,
    language: str,
    allow_ocr: bool = True,
    force_ocr: bool = False,
    on_ocr_progress: OnOcrProgress | None = None,
    on_ocr_page: OnOcrPage | None = None,
) -> ExtractedDocument:
    pages = _pdf_text_pages(path)
    should_ocr = (
        allow_ocr
        and settings.ocr_enabled
        and (force_ocr or pages_need_ocr(pages))
    )
    if should_ocr:
        if paddle_ocr_available():
            ocr_kwargs: dict[str, Any] = {
                "language": language,
                "dpi": settings.ocr_dpi,
                "settings": settings,
            }
            if on_ocr_progress is not None:
                ocr_kwargs["on_progress"] = on_ocr_progress
            if on_ocr_page is not None:
                ocr_kwargs["on_page"] = on_ocr_page
            pages = _ocr_pdf_pages_paddle(path, **ocr_kwargs)
            method = "pymupdf+paddleocr"
        else:
            logger.warning(
                "OCR needed for %s but PaddleOCR is not available",
                path,
            )
            method = "pymupdf"
    else:
        method = "pymupdf"

    return ExtractedDocument(
        pages=pages,
        page_count=len(pages),
        language=language if method != "pymupdf" else None,
        method=method,
    )


def _pdf_text_pages(path: Path) -> list[ExtractedPage]:
    pages: list[ExtractedPage] = []
    with pymupdf.open(path) as doc:
        for index in range(doc.page_count):
            page = doc.load_page(index)
            text = page.get_text("text").strip()
            pages.append(ExtractedPage(page_number=index + 1, text=text))
    return pages


def pages_need_ocr(pages: list[ExtractedPage]) -> bool:
    """True when native page text is missing or too thin for reliable filing."""
    if not pages:
        return True
    non_empty = [p for p in pages if p.text.strip()]
    if not non_empty:
        return True
    avg_chars = sum(len(p.text) for p in non_empty) / len(non_empty)
    return avg_chars < _MIN_CHARS_PER_PAGE


# Back-compat alias for older call sites / tests.
_needs_ocr = pages_need_ocr


def _extract_image(
    path: Path,
    *,
    settings: Settings,
    language: str,
    on_ocr_progress: OnOcrProgress | None = None,
    on_ocr_page: OnOcrPage | None = None,
) -> ExtractedDocument:
    text = ""
    method = "pillow"
    if settings.ocr_enabled and paddle_ocr_available():
        if settings.ocr_in_process:
            if on_ocr_progress is not None:
                on_ocr_progress(0, 1)
            text = ocr_image(path, language=language)
            method = "paddleocr"
            if on_ocr_page is not None:
                on_ocr_page(1, text)
            if on_ocr_progress is not None:
                on_ocr_progress(1, 1)
        else:
            captured: list[str] = []

            def _collect(page_number: int, page_text: str) -> None:
                captured.append(page_text)
                if on_ocr_page is not None:
                    on_ocr_page(page_number, page_text)

            try:
                done = run_ocr_subprocess(
                    mode="image",
                    path=path,
                    language=language,
                    dpi=settings.ocr_dpi,
                    timeout_seconds=settings.ocr_subprocess_timeout_seconds,
                    on_progress=on_ocr_progress,
                    on_page=_collect,
                )
            except OcrSubprocessError:
                logger.exception("OCR subprocess failed for image %s", path)
                raise
            method = done.method
            text = captured[0] if captured else ""
    else:
        with Image.open(path) as img:
            # Metadata-only fallback when OCR is unavailable.
            parts = [
                f"Image: {path.name}",
                f"Size: {img.size[0]}x{img.size[1]}",
                f"Mode: {img.mode}",
            ]
            text = "\n".join(parts)
            if not settings.ocr_enabled:
                logger.info("OCR disabled; returning image metadata for %s", path)
            elif settings.ocr_enabled:
                logger.warning(
                    "OCR enabled but PaddleOCR unavailable; metadata fallback for %s",
                    path,
                )

    page = ExtractedPage(page_number=1, text=text.strip())
    return ExtractedDocument(
        pages=[page],
        page_count=1,
        language=language if method == "paddleocr" else None,
        method=method,
    )


def _ocr_pdf_pages_paddle(
    path: Path,
    *,
    language: str,
    dpi: int | None = None,
    settings: Settings | None = None,
    on_progress: OnOcrProgress | None = None,
    on_page: OnOcrPage | None = None,
) -> list[ExtractedPage]:
    settings = settings or get_settings()
    render_dpi = int(dpi if dpi is not None else getattr(settings, "ocr_dpi", _DEFAULT_OCR_DPI))

    if settings.ocr_in_process:
        return _ocr_pdf_pages_paddle_inprocess(
            path,
            language=language,
            dpi=render_dpi,
            on_progress=on_progress,
            on_page=on_page,
        )

    pages: list[ExtractedPage] = []

    def _collect(page_number: int, text: str) -> None:
        pages.append(ExtractedPage(page_number=page_number, text=text))
        if on_page is not None:
            on_page(page_number, text)

    try:
        done = run_ocr_subprocess(
            mode="pdf",
            path=path,
            language=language,
            dpi=render_dpi,
            timeout_seconds=settings.ocr_subprocess_timeout_seconds,
            on_progress=on_progress,
            on_page=_collect,
        )
    except OcrSubprocessError:
        logger.exception("OCR subprocess failed for PDF %s", path)
        raise

    if done.page_count and len(pages) != done.page_count:
        logger.warning(
            "OCR page count mismatch for %s: events=%s done=%s",
            path,
            len(pages),
            done.page_count,
        )
    return pages


def _ocr_pdf_pages_paddle_inprocess(
    path: Path,
    *,
    language: str,
    dpi: int,
    on_progress: OnOcrProgress | None = None,
    on_page: OnOcrPage | None = None,
) -> list[ExtractedPage]:
    pages: list[ExtractedPage] = []
    with pymupdf.open(path) as doc:
        total = doc.page_count
        if on_progress is not None:
            on_progress(0, total)
        for index in range(total):
            page = doc.load_page(index)
            pix = page.get_pixmap(dpi=dpi, alpha=False)
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp_path = Path(tmp.name)
            try:
                pix.save(str(tmp_path))
                del pix
                text = ocr_image(tmp_path, language=language)
            finally:
                tmp_path.unlink(missing_ok=True)
            page_number = index + 1
            pages.append(ExtractedPage(page_number=page_number, text=text))
            if on_page is not None:
                on_page(page_number, text)
            if on_progress is not None:
                on_progress(page_number, total)
    return pages


def detect_language_hint(text: str) -> str | None:
    """Return a coarse Folium language hint from extracted text.

    Codes match OCR_LANGUAGE / legacy Tesseract-style values and are mapped to
    PaddleOCR langs by ``folium.ocr.paddle_engine.map_ocr_language``.
    """
    sample = text[:4000]
    if not sample.strip():
        return None
    if re.search(r"[\u0400-\u04FF]", sample):
        return "rus"
    if re.search(r"[\u4E00-\u9FFF]", sample):
        return "chi_sim"
    if re.search(r"[\u0600-\u06FF]", sample):
        return "ara"
    return "eng"
